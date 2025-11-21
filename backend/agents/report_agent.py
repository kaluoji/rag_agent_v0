# agents/report_agent.py - Versión corregida
from __future__ import annotations as _annotations

import os
import datetime
import logging
import re
from typing import Optional, Dict, Any, List

from pydantic import BaseModel
from pydantic_ai import Agent, RunContext
from pydantic_ai.models.openai import OpenAIModel
from openai import AsyncOpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.config import settings

# Configuración de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Inicialización del modelo
llm = settings.llm_model
model = OpenAIModel(llm)

report_system_prompt = """
Eres un experto en análisis regulatorio que genera reportes profesionales diferenciados por sección.

PRINCIPIOS CLAVE:
1. DIFERENCIACIÓN: Cada sección debe tener un propósito único y contenido específico
2. PROGRESIÓN: Las secciones deben complementarse, no repetirse
3. ESPECIFICIDAD: Usar solo información concreta de la base de conocimiento
4. ACCIONABILIDAD: Las recomendaciones deben ser implementables

FLUJO LÓGICO:
- Executive Summary: Vista de alto nivel (QUÉ)
- Alcance: Perímetro del análisis (DÓNDE y CUÁNDO)  
- Findings: Hallazgos específicos (QUÉ se encontró)
- Conclusiones: Interpretación y acciones (QUÉ hacer)

PROHIBIDO:
- Repetir la misma información en múltiples secciones
- Inventar artículos o disposiciones no documentadas
- Usar lenguaje genérico sin base en la documentación
- Generar recomendaciones sin fundamento específico

Template fijo: "Template_Regulatory_Report_AgentIA_v0.docx"
"""

class ReportDeps(BaseModel):
    output_folder: str = "output/reports"
    template_folder: str = "agents/templates"
    openai_client: AsyncOpenAI

    class Config:
        arbitrary_types_allowed = True

class ReportResult(BaseModel):
    file_path: str
    report_type: str = "word"
    message: str = "Reporte generado exitosamente"
    sections_filled: List[str] = []

# Mapeo de placeholders del template
TEMPLATE_PLACEHOLDERS = {
    "{{RESUMEN_EJECUTIVO}}": "executive_summary",
    "{{ALCANCE_ANALISIS}}": "scope", 
    "{{ANALISIS}}": "findings",
    "{{ANALISIS_SECTORIAL}}": "sector_impact_analysis",
    "{{RECOMENDACIONES}}": "conclusions_recommendations"

}

report_agent = Agent(
    model=model,
    system_prompt=report_system_prompt,
    deps_type=ReportDeps,
    output_type=ReportResult,
    retries=2
)

def find_placeholders_in_document(doc_path: str) -> List[str]:
    """
    Encuentra todos los placeholders en un documento Word.
    """
    doc = Document(doc_path)
    placeholders = []
    
    # Buscar en párrafos
    for paragraph in doc.paragraphs:
        text = paragraph.text
        found = re.findall(r'\{\{[^}]+\}\}', text)
        placeholders.extend(found)
    
    # Buscar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    found = re.findall(r'\{\{[^}]+\}\}', text)
                    placeholders.extend(found)

    # 🔹 NUEVO: buscar también en TODO el XML (incluye cuadros de texto, etc.)
    xml = doc._element.xml
    found_xml = re.findall(r'\{\{[^}]+\}\}', xml)
    placeholders.extend(found_xml)
    
    logger.info(f"Placeholders encontrados en el documento: {placeholders}")
    return list(set(placeholders))

def normalize_generated_text(text: str) -> str:
    """
    Limpia un poco el texto generado por el modelo para que Word lo justifique mejor:
    - Une líneas sueltas en un mismo párrafo.
    - Mantiene sólo los saltos de párrafo (doble salto de línea).
    """
    # Separar por párrafos (doble salto de línea)
    paragraphs = re.split(r'\n\s*\n', text)
    cleaned_paragraphs = []

    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        # Dentro de cada párrafo, sustituir saltos de línea simples por espacios
        one_line = re.sub(r'\s*\n\s*', ' ', p)
        # Colapsar espacios múltiples
        one_line = re.sub(r' {2,}', ' ', one_line)
        cleaned_paragraphs.append(one_line)

    return "\n\n".join(cleaned_paragraphs)


def replace_placeholder_in_document(doc: Document, placeholder: str, content: str):
    """
    Reemplaza un placeholder específico en todo el documento con el contenido generado.
    """
    replacements_made = 0
    
    # Reemplazar en párrafos
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            new_text = paragraph.text.replace(placeholder, content)
            paragraph.clear()
            insert_markdown(paragraph, new_text)
            replacements_made += 1
    
    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        new_text = paragraph.text.replace(placeholder, content)
                        paragraph.clear()
                        insert_markdown(paragraph, new_text)
                        replacements_made += 1

    # 🔹 NUEVO: reemplazar también en cualquier nodo de texto del XML
    if placeholder in ("{{FECHA}}", "{{LEY_ANALIZADA}}"):
        for element in doc._element.iter():
            if element.tag.endswith('}t') and element.text and placeholder in element.text:
                element.text = element.text.replace(placeholder, content)
                replacements_made += 1
    
    logger.info(f"Realizados {replacements_made} reemplazos para el placeholder '{placeholder}'")

def insert_markdown_bold(paragraph, text):
    """
    Inserta texto en un párrafo de Word interpretando **negrita** estilo Markdown.
    """
    import re
    pattern = r'\*\*(.*?)\*\*'
    last_index = 0

    for match in re.finditer(pattern, text):
        # Texto normal antes de la negrita
        if match.start() > last_index:
            paragraph.add_run(text[last_index:match.start()])

        # Texto en negrita
        bold_run = paragraph.add_run(match.group(1))
        bold_run.bold = True

        last_index = match.end()

    # Texto restante
    if last_index < len(text):
        paragraph.add_run(text[last_index:])

def insert_markdown(paragraph, text: str):
    """
    Inserta contenido interpretando Markdown básico EN LA POSICIÓN DEL PLACEHOLDER:
    - # Título 1  -> estilo 'Heading 1'
    - ## Título 2 -> estilo 'Heading 2'
    - ### Título 3 (o más #) -> estilo 'Heading 3'
    - - item      -> lista con viñeta ('List Bullet')
    - 1. item     -> lista numerada ('List Number')
    - **texto**   -> negrita en runs
    Cada línea se convierte en un párrafo nuevo o en el párrafo original.
    """
    import re

    # Separar en líneas y limpiar espacios
    # Unificar líneas normales dentro de un mismo párrafo
    raw_lines = text.split("\n")
    lines = []
    buffer = ""

    for l in raw_lines:
        l = l.strip()

        # Si es un encabezado o lista, se corta el buffer
        if re.match(r"^(#{1,6})\s+", l) or l.startswith("- ") or re.match(r"^\d+\.\s+", l):
            if buffer:
                lines.append(buffer.strip())
                buffer = ""
            lines.append(l)
        else:
            # Línea normal → acumularla en buffer
            if l:
                buffer += " " + l

    if buffer:
        lines.append(buffer.strip())


    # Eliminar líneas vacías al principio y al final
    while lines and lines[0] == "":
        lines.pop(0)
    while lines and lines[-1] == "":
        lines.pop()

    if not lines:
        return

    current_paragraph = paragraph  # empezamos usando el párrafo del placeholder

    # Procesamos en orden inverso para que insert_paragraph_before mantenga el orden final correcto
    first = True
    for line in reversed(lines):
        line = line.strip()

        if line == "":
            # línea en blanco -> párrafo vacío antes
            new_p = current_paragraph.insert_paragraph_before()
            current_paragraph = new_p
            continue

        # 🔹 Saltar separadores tipo '---', '***', '___'
        if line in ("---", "***", "___"):
            new_p = current_paragraph.insert_paragraph_before()
            current_paragraph = new_p
            first = False
            continue

        # Detectar tipo de línea
        style = None
        content = line

        # 🔹 Encabezados Markdown (#, ##, ###, etc.)
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            hashes, rest = m.groups()
            level = len(hashes)
            content = rest.strip()
            if level == 1:
                style = "Heading 1"
            elif level == 2:
                style = "Heading 2"
            else:
                style = "Heading 3"
        # Lista numerada (1. , 2. , etc.)
        elif re.match(r"^\d+\.\s+", line):
            style = "List Number"
            content = re.sub(r"^\d+\.\s+", "", line).strip()
        # Lista con viñeta
        elif line.startswith("- "):
            style = "List Bullet"
            content = line[2:].strip()

        # Elegir en qué párrafo escribir
        if first:
            # La última línea se escribe en el párrafo original
            target_p = current_paragraph
            first = False
        else:
            # Las demás se insertan antes, de abajo hacia arriba
            target_p = current_paragraph.insert_paragraph_before()
            current_paragraph = target_p

        # Aplicar estilo si procede
        if style:
            try:
                target_p.style = style
            except Exception:
                # Si el estilo no existe en la plantilla, lo ignoramos y seguimos
                pass

        # Rellenar el párrafo con negritas interpretando **texto**
        insert_markdown_bold(target_p, content)




@report_agent.tool
async def generate_report_from_template(
    ctx: RunContext[ReportDeps], 
    analysis_data: str, 
    template_name: str = "Template_Regulatory_Report_AgentIA_v0.docx",
    regulation_name: str = "Normativa Analizada",
    output_filename: str = None
) -> ReportResult:
    """
    Genera un reporte usando un template de Word existente.
    """
    try:
        # Construir rutas
        template_path = os.path.join(ctx.deps.template_folder, template_name)
        
        if not os.path.exists(template_path):
            logger.error(f"Template no encontrado en: {template_path}")
            # Listar archivos disponibles para debugging
            if os.path.exists(ctx.deps.template_folder):
                available_files = os.listdir(ctx.deps.template_folder)
                logger.info(f"Archivos disponibles en template_folder: {available_files}")
            raise FileNotFoundError(f"Template no encontrado: {template_path}")
        
        # Cargar el template
        doc = Document(template_path)
        logger.info(f"Template cargado exitosamente desde: {template_path}")
        
        # Encontrar placeholders en el documento
        placeholders = find_placeholders_in_document(template_path)
        logger.info(f"Placeholders encontrados: {placeholders}")
        
        if not placeholders:
            logger.warning("No se encontraron placeholders en el documento template")
        
        # Reemplazar fecha y nombre de regulación
        current_date = datetime.datetime.now().strftime("%d/%m/%Y")
        replace_placeholder_in_document(doc, "{{FECHA}}", current_date)
        replace_placeholder_in_document(doc, "{{LEY_ANALIZADA}}", regulation_name)
        
        # Generar contenido para cada sección
        sections_filled = []
        
        for placeholder in placeholders:
            logger.info(f"Procesando placeholder: {placeholder}")
            
            # Buscar el placeholder en nuestro mapeo
            section_key = None
            for mapped_placeholder, key in TEMPLATE_PLACEHOLDERS.items():
                if placeholder.strip() == mapped_placeholder.strip():
                    section_key = key
                    break
            
            if section_key:
                logger.info(f"Generando contenido para sección: {section_key}")
                content = await generate_section_content(
                    ctx.deps.openai_client, 
                    section_key, 
                    analysis_data, 
                    regulation_name
                )

                content = normalize_generated_text(content)

                replace_placeholder_in_document(doc, placeholder, content)

                sections_filled.append(section_key)
                logger.info(f"Sección completada: {section_key}")
            else:
                logger.warning(f"Placeholder no reconocido: {placeholder}")
        
        # Crear directorio de salida si no existe
        if not os.path.exists(ctx.deps.output_folder):
            os.makedirs(ctx.deps.output_folder)
            logger.info(f"Directorio creado: {ctx.deps.output_folder}")
        
        # Generar nombre de archivo
        if output_filename is None:
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_regulation_name = re.sub(r'[^\w\s-]', '', regulation_name).strip()
            safe_regulation_name = re.sub(r'[-\s]+', '_', safe_regulation_name)
            output_filename = f"Reporte_Normativo_{safe_regulation_name}_{timestamp}.docx"
        
        output_path = os.path.join(ctx.deps.output_folder, output_filename)
        doc.save(output_path)
        
        logger.info(f"Reporte generado exitosamente: {output_path}")
        
        return ReportResult(
            file_path=output_path,
            report_type="word",
            message=f"Reporte generado usando template {template_name}. Secciones completadas: {len(sections_filled)}",
            sections_filled=sections_filled
        )
        
    except Exception as e:
        logger.error(f"Error generando reporte desde template: {e}")
        raise e

async def generate_section_content(
    openai_client: AsyncOpenAI, 
    section_key: str, 
    analysis_data: str, 
    regulation_name: str
) -> str:
    """
    Genera contenido específico para una sección del reporte usando EXCLUSIVAMENTE
    la información ya recuperada de Supabase por el agente de compliance.
    """
    
    # MEJORA: Analizar si analysis_data contiene información de Supabase
    has_supabase_data = "Content:" in analysis_data or "title" in analysis_data.lower() or "artículo" in analysis_data.lower()
    
    section_prompts = {
    "executive_summary": f"""
INFORMACIÓN DE LA BASE DE CONOCIMIENTO:
{analysis_data}

OBJETIVO:
Ofrecer una visión clara, estratégica y ejecutiva de {regulation_name} para permitir que directivos y áreas de cumplimiento comprendan rápidamente qué cambia, por qué importa y dónde deben enfocar su atención inmediata.

TAREA:
Redacta un RESUMEN EJECUTIVO (máx. 1000 palabras) sobre {regulation_name}, sintetizando el propósito de la norma, sus implicaciones y el impacto sectorial.

ESTRUCTURA:
1. Contexto y propósito de la regulación
2. Cambios principales u obligaciones clave extraídas del texto
3. Impacto estratégico y operativo para el sector afectado
4. Conclusión ejecutiva con foco inmediato para áreas jurídicas y de compliance

REGLAS:
- Evitar repeticiones con las secciones de Findings o Recomendaciones
- Tono ejecutivo, orientado a toma de decisiones
- Sin adornos ni lenguaje genérico
- Crear una transición natural hacia “Alcance del análisis”

FORMATO:
- Puedes usar Markdown para estructurar la información.
- Usa **negrita** con Markdown.
- Usa solo encabezados de nivel 2 o inferior (##, ###) para subapartados internos.
- NO incluyas un título tipo “Resumen ejecutivo” ni el nombre de la norma; el título de la sección ya está en el documento. Empieza directamente con el contenido.
- Usa listas (- o 1.) cuando sea más claro que un párrafo.

PROHIBIDO:
- Agregar información no presente en el contexto
- Usar lenguaje genérico sin métricas/plazos concretos
- Repetir contenido que irá en otras secciones

ESTILO: Prosa directa, datos específicos, sin listas.
""",


    "scope": f"""
INFORMACIÓN DE LA BASE DE CONOCIMIENTO:
{analysis_data}

OBJETIVO:
Delimitar claramente qué parte de la regulación se ha analizado (sin incluir el número de los artículos), qué queda fuera, qué áreas funcionales se ven afectadas y bajo qué criterios se ha estructurado el análisis, para asegurar trazabilidad y rigor metodológico.

TAREA:
Define el ALCANCE DEL ANÁLISIS para {regulation_name}.

ESTRUCTURA:
1. Títulos y disposiciones examinadas
2. Áreas organizativas del sector afectadas según el contenido disponible
3. Periodo temporal o contexto de aplicación
4. Exclusiones del análisis y limitaciones documentales
5. Metodología utilizada y criterios de priorización

FORMATO:
- Puedes usar Markdown para estructurar la información.
- Usa **negrita** con Markdown para resaltar conceptos clave.
- Usa solo encabezados de nivel 2 o inferior (##, ###) para las subsecciones.
- NO repitas el título “Alcance del análisis” ni variantes; el título ya está en el template. Empieza directamente con el punto 1.
- Usa listas (- o 1.) cuando sea más claro que un párrafo.


REGLAS:
- Nada de conclusiones ni impactos (eso va en Findings)
- Redacción técnica, concisa y trazable
- Mencionar solo lo presente en el análisis_data
- Preparar la transición a “Análisis normativo / Findings”
""",

    "findings": f"""
INFORMACIÓN DE LA BASE DE CONOCIMIENTO:
{analysis_data}

OBJETIVO:
Identificar, clasificar y explicar las obligaciones normativas relevantes, así como su impacto operativo y su nivel de riesgo para las organizaciones del sector, proporcionando la base para recomendaciones accionables.

TAREA:
Desarrolla el ANÁLISIS NORMATIVO / FINDINGS de {regulation_name}, agrupando las obligaciones por nivel de impacto y explicando su relevancia operativa.

ESTRUCTURA:
1. Obligaciones críticas (impacto alto / cumplimiento inmediato)
   - Artículo
   - Obligación
   - Riesgo o sanción por incumplimiento
   - Área o proceso impactado

2. Obligaciones relevantes (impacto medio / seguimiento continuo)
   - Requisito
   - Implicación operativa o de gobernanza

3. Aspectos informativos o de contexto (impacto bajo)
   - Conceptos clave que guían la norma

4. Brechas, vacíos o ambigüedades detectadas

REGLAS:
- Lenguaje claro y cuantificable
- Evitar generalidades jurídicas
- Conectar cada hallazgo con su posible traducción operativa

FORMATO:
- Puedes usar Markdown para estructurar la información.
- Usa **negrita** para los identificadores de hallazgos (C1, C2, M1, etc.) y para conceptos clave como plazos y tipos de reporte.
- Usa encabezados (###) para cada hallazgo (por ejemplo: "### C1. Obligación...").
- NO generes un encabezado general del tipo “ANÁLISIS NORMATIVO / FINDINGS”; ese título ya existe en el documento.
- Usa listas (- o 1.) cuando sea más claro que un párrafo.


""",

"sector_impact_analysis": f"""
INFORMACIÓN DE LA BASE DE CONOCIMIENTO:
{analysis_data}

OBJETIVO DE LA SECCIÓN:
Analizar cómo los requisitos de {regulation_name} afectan de manera específica al sector evaluado, describiendo cambios operativos, riesgos, retos y oportunidades. El objetivo es permitir que las áreas de jurídica y compliance traduzcan los efectos de la norma en impactos reales sobre procesos, modelos de negocio y funciones clave.

TAREA:
Redacta un ANÁLISIS DE IMPACTO SECTORIAL basado EXCLUSIVAMENTE en la información disponible, interpretando de forma profesional y sin añadir conocimiento externo qué efectos puede generar la normativa sobre las actividades típicas del sector (por ejemplo, asegurador, bancario, tecnológico o equivalente, según se deduzca del contenido).

ESTRUCTURA OBLIGATORIA (no modificar):
1. Procesos o actividades del sector directamente afectados  
2. Cambios operativos, tecnológicos o documentales derivados de la norma  
3. Riesgos sectoriales emergentes (operativos, sancionadores, reputacionales)  
4. Oportunidades de mejora o ventajas competitivas derivadas de la adaptación  
5. Implicaciones para modelos de gobierno, control interno y reporting  

REGLAS ESTRICTAS:
- NO añadir información externa ni interpretaciones que no estén respaldadas por el análisis previo.  
- Basar cada punto en disposiciones, artículos o principios citados en {analysis_data}.  
- Lenguaje profesional, claro y orientado a toma de decisiones.  
- No repetir los ‘findings’; esta sección debe interpretar SU IMPACTO en el sector.    
- Mantener conexión explícita entre los impactos y los artículos u obligaciones relevantes.  

FORMATO:
- Puedes usar Markdown para estructurar la información.
- Usa **negrita** con Markdown para conceptos clave.
- Usa solo encabezados de nivel 2 o inferior (##, ###) para ordenar los 5 apartados, si lo consideras útil.
- NO añadas un título general como “Análisis de Impacto Sectorial”; ya existe en el template.
- Usa listas (- o 1.) cuando sea más claro que un párrafo.


OBJETIVO FINAL DEL TEXTO:
Brindar una visión estratégica del impacto real de la regulación sobre el sector, sirviendo como puente entre los hallazgos normativos y las recomendaciones operativas del informe.
""",

    "conclusions_recommendations": f"""
INFORMACIÓN DE LA BASE DE CONOCIMIENTO:
{analysis_data}

OBJETIVO:
Transformar los hallazgos normativos en acciones claras, priorizadas y ejecutables por áreas jurídicas, técnicas y de negocio, permitiendo avanzar hacia el cumplimiento efectivo y la preparación ante auditorías.

TAREA:
Genera RECOMENDACIONES ESTRATÉGICAS y ACCIONABLES basadas en los hallazgos del análisis.

ESTRUCTURA:

1. Acciones inmediatas (0–30 días)
   - Acción específica
   - Responsable sugerido
   - Entregable verificable
   - Hallazgo/artículo asociado

2. Implementación a medio plazo (1–6 meses)
   - Proyecto o mejora
   - Recursos requeridos
   - Resultado esperado
   - Hallazgo/artículo asociado

3. Monitoreo y mejora continua
   - Indicadores o KPIs
   - Frecuencia
   - Responsable

4. Impacto sectorial consolidado
   - Procesos del sector más afectados
   - Riesgos o oportunidades emergentes

REGLAS:
- Todo debe vincularse a hallazgos concretos
- Tono consultivo, claro y orientado al cumplimiento
- Evitar abstracciones o recomendaciones genéricas
- Finalizar con una visión sintética del roadmap normativo

FORMATO:
- Usa **negrita** para resaltar acciones clave, responsables y plazos.
- Usa encabezados (##, ###) para dividir bloques (Acciones inmediatas, Medio plazo, etc.).
- NO generes un encabezado general “Recomendaciones” o similar; el título de la sección ya está en el documento.
- Usa listas (- o 1.) para detallar acciones y KPIs.


"""
}
    
    prompt = section_prompts.get(section_key, f"Genera contenido para la sección {section_key} basado EXCLUSIVAMENTE en: {analysis_data}")
    
    try:
        completion = await openai_client.chat.completions.create(
            model=settings.llm_model,
            temperature=0.2,  # Temperatura baja para ser más fiel a los datos
            max_tokens=5000,
            messages=[
                {
                    "role": "system", 
                    "content": """Eres un experto en análisis regulatorio de sector asegurador que trabaja EXCLUSIVAMENTE con información específica de bases de conocimiento.

INSTRUCCIONES CRÍTICAS:
1. Usa SOLO la información proporcionada en el contexto.
2. No agregues conocimiento general ni información externa.
3. Cita específicamente artículos, secciones y disposiciones encontradas, si existen.
4. Si la información es limitada o faltante, dilo explícitamente y señala la sección “Alcance” como referencia de límites.
5. Permite Markdown básico: negritas (**texto**), encabezados de nivel 2 o inferior (##, ###) y listas (-, 1.); NO uses encabezados de nivel 1 (#) porque el título principal de cada sección ya está en el template.
6. Mantén fidelidad absoluta a la documentación proporcionada.
7. Garantiza coherencia interseccional: las conclusiones deben referenciar hallazgos; los hallazgos deben alinearse con el alcance; el resumen ejecutivo debe mencionar dónde ampliar.
8. Emplea conectores y frases puente para mejorar la fluidez (por ejemplo: “En consecuencia…”, “De acuerdo con…”, “Como se detalla en Hallazgos…”).
9. Reutiliza terminología y definiciones tal como aparecen en el contexto; si hay siglas, estandarízalas y utilízalas consistentemente.

PROHIBIDO:
- Inventar artículos o disposiciones no mencionadas.
- Agregar información de conocimiento general.
- Hacer suposiciones no respaldadas por la documentación.

REGLAS DE TRAZABILIDAD:
- Asigna identificadores a los hallazgos para permitir referencias cruzadas:
  • Críticos: C1, C2
  • Medios: M1, M2
  • Informativos: I1, I2
- En las recomendaciones, referencia SIEMPRE los identificadores de hallazgos (por ejemplo, "Relacionado con: C1").

VALIDACIÓN DE COHERENCIA INTERSECCIONAL:
- Si mencionas un hallazgo [ID] en Resumen o Recomendaciones, DEBE existir en Findings
- Si citas un artículo en Recomendaciones, DEBE haber sido mencionado en Findings o Alcance
- Los plazos mencionados deben ser consistentes en todas las secciones
- La terminología técnica debe usarse de forma uniforme

"""
                },
                {"role": "user", "content": prompt}
            ]
        )
        
        content = completion.choices[0].message.content.strip()
        
        
        # VERIFICACIÓN: Asegurar que se está usando información de Supabase
        if has_supabase_data:
            logger.info(f"Sección {section_key} generada usando información de Supabase")
        else:
            logger.warning(f"Sección {section_key} generada con información limitada - verificar recuperación de Supabase")
        
        return content
        
    except Exception as e:
        logger.error(f"Error generando contenido para sección {section_key}: {e}")
        return f"[Error generando contenido para {section_key}]"

async def process_report_query(
    query: str, 
    analysis_data: str, 
    deps: ReportDeps,
    template_name: str = "Template_Regulatory_Report_AgentIA_v0.docx",
    regulation_name: str = None
) -> ReportResult:
    """
    Procesa una consulta para generar un reporte usando un template específico.
    """
    logger.info(f"Procesando consulta con template: {query[:100]}...")
    
    # Extraer nombre de regulación si no se proporciona
    if regulation_name is None:
        try:
            completion = await deps.openai_client.chat.completions.create(
                model=settings.llm_model,
                temperature=0.1,
                messages=[
                    {
                        "role": "system", 
                        "content": "Extrae el nombre de la regulación o ley mencionada en la consulta y su denominación comunmente conocida. Responde solo con el nombre y la denominación comunmente conocida."
                    },
                    {"role": "user", "content": query}
                ]
            )
            regulation_name = completion.choices[0].message.content.strip()
        except:
            regulation_name = "Regulación Analizada"
    
    # CORRECCIÓN: Crear un RunContext simulado usando los datos necesarios
    try:
        # Llamar directamente a la función generate_report_from_template
        # creando un contexto simulado
        class MockRunContext:
            def __init__(self, deps):
                self.deps = deps
        
        mock_ctx = MockRunContext(deps)
        
        return await generate_report_from_template(
            mock_ctx,
            analysis_data=analysis_data,
            template_name=template_name,
            regulation_name=regulation_name
        )
        
    except Exception as e:
        logger.error(f"Error procesando consulta con template: {e}")
        return ReportResult(
            file_path="",
            message=f"Error generando reporte con template: {str(e)}"
        )